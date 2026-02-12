from docx import Document

def read_docx(file_path):
    try:
        document = Document(file_path)
        full_text = []
        print(f"Number of paragraphs: {len(document.paragraphs)}")
        for paragraph in document.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        if not full_text:
            print("No text found in paragraphs. Checking for text in tables...")
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

        return "\n".join(full_text) if full_text else "No readable content found."
    except Exception as e:
        return f"Error reading DOCX file: {e}"

if __name__ == "__main__":
    file_path = "data/demo_wills/will_sample_1.docx"
    content = read_docx(file_path)
    print(content)
