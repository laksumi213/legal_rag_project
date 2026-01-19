# src/legal_system/ui/label_generator.py

import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 拠点ごとの住所定義
BRANCH_ADDRESSES = {
    "東京": "〒103-0028\n東京都中央区八重洲一丁目7-20\n八重洲口会館2階",
    "横浜": "〒220-0011\n神奈川県横浜市西区高島2-19-12\nスカイビル",
    "新宿": "〒163-0000\n東京都新宿区西新宿..." 
}

def get_branch_address(branch_name: str) -> str:
    for key in BRANCH_ADDRESSES:
        if key in branch_name:
            return BRANCH_ADDRESSES[key]
    return BRANCH_ADDRESSES["東京"]

def _set_font_style(run, size_pt=12, is_bold=False):
    """フォントスタイル（MS明朝）を一括適用するヘルパー"""
    run.font.name = "MS Mincho"
    run.font.size = Pt(size_pt)
    run.font.bold = is_bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'MS Mincho')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'MS Mincho')

def _get_cell_width(cell):
    """セルの幅をXMLから取得する（Twips単位）"""
    try:
        tcW = cell._tc.tcPr.tcW
        if tcW.type == 'dxa':
            return int(tcW.w)
        elif tcW.type == 'pct':
            return 0
    except:
        pass
    return 0

def _get_valid_cells(table):
    """幅が極端に狭いセル（余白用スペーサー）を除外する"""
    valid_cells = []
    all_widths = []
    flat_cells = []
    
    for row in table.rows:
        for cell in row.cells:
            w = _get_cell_width(cell)
            all_widths.append(w)
            flat_cells.append(cell)
            
    if not all_widths:
        return flat_cells
        
    max_width = max(all_widths)
    threshold = max_width * 0.5 
    
    for cell, w in zip(flat_cells, all_widths):
        if w == 0 or w > threshold:
            valid_cells.append(cell)
            
    return valid_cells

def _write_cell_content(cell, text_info, is_sender=False):
    """セルにテキストを書き込む（余白調整付き）"""
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- 修正: 上部の余白 (1行あける) ---
    p.add_run("\n")

    # 左余白用のスペース
    pad = " "

    # 1. 郵便番号 (12pt)
    if text_info.get("zip_code"):
        # 半角スペース + 〒...
        run_zip = p.add_run(f"{pad}〒{text_info['zip_code']}\n")
        _set_font_style(run_zip, size_pt=12)

    # 2. 住所 (12pt)
    addr = text_info.get("address", "")
    # 半角スペース + 住所...
    run_addr = p.add_run(f"{pad}{addr}\n\n")
    _set_font_style(run_addr, size_pt=12)

    # 3. 氏名 (14pt Bold)
    name_str = text_info.get("name", "")
    honor = text_info.get("honorific", "")
    
    # 半角スペース + 氏名...
    run_name = p.add_run(f"{pad}{name_str} {honor}")
    _set_font_style(run_name, size_pt=14, is_bold=True)
    
    # 4. TEL (12pt)
    if text_info.get("tel"):
        # 改行 + 半角スペース + TEL...
        run_tel = p.add_run(f"\n{pad}TEL: {text_info['tel']}")
        _set_font_style(run_tel, size_pt=12)

def generate_advanced_label(
    template_bytes: bytes, 
    print_list: list, 
    start_position: int = 1
) -> io.BytesIO:
    """高度なラベル生成関数"""
    target_stream = io.BytesIO()
    target_stream.write(template_bytes)
    target_stream.seek(0)
    
    doc = Document(target_stream)
    
    if not doc.tables:
        raise ValueError("Wordファイルに表(テーブル)が見つかりません。")

    table = doc.tables[0]
    valid_cells = _get_valid_cells(table)
    
    current_idx = start_position - 1
    
    for data in print_list:
        if current_idx >= len(valid_cells):
            break
            
        target_cell = valid_cells[current_idx]
        is_sender = data.get("type") == "sender"
        
        _write_cell_content(target_cell, data, is_sender)
        
        current_idx += 1

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream