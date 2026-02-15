from docx import Document
from io import BytesIO
import PyPDF2
import os
from pathlib import Path

class DocumentLoader:
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.file_bytes = None
        if self.file_path.exists():
            with open(self.file_path, "rb") as f:
                self.file_bytes = f.read()

    def load_and_extract_text(self) -> str:
        if not self.file_bytes:
            return ""

        file_extension = self.file_path.suffix.lower()

        if file_extension == ".docx":
            return self._extract_text_from_docx()
        elif file_extension == ".pdf":
            return self._extract_text_from_pdf()
        elif file_extension == ".txt":
            return self._extract_text_from_txt()
        else:
            # サポートされていないファイルタイプの場合、空文字列を返すか、エラーをログに記録
            return ""

    def _extract_text_from_docx(self) -> str:
        try:
            doc = Document(BytesIO(self.file_bytes))
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip(): # 空の段落はスキップ
                    full_text.append(paragraph.text)
            # テーブル内のテキストもチェック
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from DOCX {self.file_path}: {e}") # エラーをログに記録
            return ""

    def _extract_text_from_pdf(self) -> str:
        try:
            reader = PyPDF2.PdfReader(BytesIO(self.file_bytes))
            full_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text: # 空のページはスキップ
                    full_text.append(text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from PDF {self.file_path}: {e}") # エラーをログに記録
            return ""

    def _extract_text_from_txt(self) -> str:
        try:
            return self.file_bytes.decode("utf-8")
        except Exception as e:
            print(f"Error extracting text from TXT {self.file_path}: {e}") # エラーをログに記録
            return ""
