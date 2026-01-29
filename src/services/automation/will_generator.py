# src/services/automation/will_generator.py

import pandas as pd
import numpy as np
import base64
import re  # 正規表現モジュールを確実にインポート
from io import BytesIO
from typing import List, Tuple, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.prompts import ChatPromptTemplate, HumanMessagePromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from PIL import Image, ImageOps, ImageChops
from pypdf import PdfReader
from pdf2image import convert_from_bytes

from src.legal_system.core.ai_factory import AIFactory
from src.legal_system.core.schemas import WillDraftStructure

class WillDraftGenerator:
    def __init__(self):
        self.llm = AIFactory.get_llm(mode="cloud", temperature=0.0)

    def generate_draft(self, excel_file: BytesIO, template_file: BytesIO, registry_files: List[Any] = None) -> Tuple[BytesIO, Optional[BytesIO], WillDraftStructure, str]:
        """
        遺言書生成のメイン処理
        """
        # 1. Excel解析
        excel_file.seek(0)
        if hasattr(excel_file, 'name') and excel_file.name.endswith('.xlsx'):
            df = pd.read_excel(excel_file)
        else:
            df = pd.read_csv(excel_file)

        # データ補完
        df = df.replace(r'^\s*$', np.nan, regex=True).ffill()
        if 'No' in df.columns:
            df = df.dropna(subset=['No'])

        if df.empty:
            raise ValueError("有効なデータ行がありません。")

        csv_text = df.fillna("").to_csv(index=False)

        # 2. 登記情報の処理 (AIによるフォーマット変換)
        registry_data = self._process_registry_files(registry_files)

        # 3. AI推論 (条文構成)
        draft_data = self._invoke_ai_reasoning(csv_text)

        # 4. 遺言書Word生成 (本体 + テキストデータ)
        template_file.seek(0)
        safe_template = BytesIO(template_file.read())
        # ここで登記情報のテキストを渡す
        will_doc = self._create_will_document(safe_template, draft_data, registry_data.get("text", ""))
        
        will_stream = BytesIO()
        will_doc.save(will_stream)
        will_stream.seek(0)

        # 5. 登記情報Word生成 (別冊・画像のみ)
        registry_stream = None
        if registry_data.get("images"):
            reg_doc = self._create_registry_document(registry_data)
            registry_stream = BytesIO()
            reg_doc.save(registry_stream)
            registry_stream.seek(0)
        
        return will_stream, registry_stream, draft_data, csv_text

    def _process_registry_files(self, files: List[Any]) -> Dict[str, Any]:
        """
        登記情報(PDF/画像)を処理する
        - 画像への変換 & 余白トリミング
        - Gemini Visionによる指定フォーマットでのテキスト化
        """
        processed = {"images": [], "text": ""}
        if not files:
            return processed

        all_images_for_ai = [] # テキスト解析用に全ての画像をリスト化
        
        for f in files:
            f.seek(0)
            file_bytes = f.read()
            file_name = getattr(f, "name", "unknown")

            # PDFの場合
            if file_name.lower().endswith(".pdf") or f.type == "application/pdf":
                try:
                    # 画像変換 (200dpi)
                    pil_images = convert_from_bytes(file_bytes, dpi=200)
                    for img in pil_images:
                        # 1. 別冊用画像 (トリミング済)
                        trimmed = self._trim_whitespace(img)
                        buf = BytesIO()
                        trimmed.save(buf, format="JPEG")
                        processed["images"].append(BytesIO(buf.getvalue()))
                        
                        # 2. AI解析用画像
                        ai_buf = BytesIO()
                        img.convert("RGB").save(ai_buf, format="JPEG")
                        all_images_for_ai.append(BytesIO(ai_buf.getvalue()))

                except Exception as e:
                    print(f"PDF process error: {e}")

            # 画像の場合
            else:
                try:
                    img = Image.open(BytesIO(file_bytes))
                    
                    # 1. 別冊用
                    trimmed = self._trim_whitespace(img)
                    buf = BytesIO()
                    trimmed.save(buf, format="JPEG")
                    processed["images"].append(BytesIO(buf.getvalue()))

                    # 2. AI解析用
                    ai_buf = BytesIO()
                    img.convert("RGB").save(ai_buf, format="JPEG")
                    all_images_for_ai.append(BytesIO(ai_buf.getvalue()))

                except Exception as e:
                    print(f"Image load error: {e}")

        # --- AIによるテキスト化 (Gemini Vision) ---
        if all_images_for_ai:
            processed["text"] = self._analyze_registry_images_with_ai(all_images_for_ai)
        
        return processed

    def _analyze_registry_images_with_ai(self, image_buffers: List[BytesIO]) -> str:
        """登記情報の画像をAIに読み取らせて、指定フォーマットのテキストに変換する"""
        prompt = """
        提供された不動産登記情報の画像を読み取り、公証人が遺言書作成に使用するためのテキストデータを作成してください。
        
        # 【重要】生成ルール
        1. **所在の結合**:
           - 建物の「所在」欄にある「市区町村名」と、その下（または横）にある「地番（または家屋番号の番地部分）」を**必ず1行に結合**してください。
           - 画像上で改行されていても、出力時は全角スペースでつないで1行にしてください。
           - 例:
             [画像]
               四街道市旭ケ丘五丁目
               １５２０番２３６
             [出力]
               所在　四街道市旭ケ丘五丁目　１５２０番２３６

        2. **床面積の改行禁止**:
           - 建物が複数階ある場合でも、**絶対に改行せず**、全角スペースで区切って一行にまとめてください。
           - 例: 1階 79.08　2階 52.58㎡

        3. **マンション判定**: 
           - 文中に「一棟の建物の表示」および「敷地権」という文言が含まれる場合のみ「マンション（区分所有建物）」として扱ってください。それ以外は「土地」または「建物」です。

        4. **持分（シェア）の特定**:
           - 持分は通常、所有者氏名の直上（または直近）に記載されています（例：「持分２分の１」）。
           - 単独所有で持分の記載がない場合は空欄にしてください。（「1/1」と補完しないでください）

        5. **文字の正規化**: 
           - 氏名や地名に含まれる空白（全角・半角スペース）はすべて削除して認識してください。
           - 「ヶ」「ケ」の表記揺れは、登記簿の記載通りにしてください。

        # 出力フォーマット例
        物件ごとに（１）、（２）...と連番を振ってください。

        【土地の場合】
        （Ｎ）　土地
        　所在　■■市■■区■■■　■■番地■
        　地番　■番■
        　地目　■■
        　地積　■.■㎡
        　持分　■分の■（※記載がある場合のみ）

        【建物の場合】
        （Ｎ）　建物
        　所在　■■市■■区■■■　■■番地■
        　家屋番号　■番■
        　種類　■■
        　構造　■■
        　床面積　1階 ■.■　2階 ■.■㎡
        　持分　■分の■（※記載がある場合のみ）
        """
        
        content = [{"type": "text", "text": prompt}]
        
        for img_buf in image_buffers:
            img_buf.seek(0)
            b64_data = base64.b64encode(img_buf.read()).decode("utf-8")
            content.append({
                "type": "image_url",
                "image_url": f"data:image/jpeg;base64,{b64_data}"
            })
            
        msg = HumanMessage(content=content)
        
        try:
            res = self.llm.invoke([msg])
            raw_text = res.content
            
            # ★追加: Python側での強力な後処理（強制結合）
            return self._post_process_ai_text(raw_text)

        except Exception as e:
            return f"※AI解析エラー: {e}"

    def _post_process_ai_text(self, text: str) -> str:
        """
        AIの出力テキストに対して、正規表現を使って強制的に行を結合する。
        """
        lines = text.split('\n')
        processed_lines = []
        
        skip_next = False
        
        for i in range(len(lines)):
            if skip_next:
                skip_next = False
                continue
            
            line = lines[i].strip()
            
            # 末尾の行ならそのまま追加
            if i == len(lines) - 1:
                processed_lines.append(lines[i])
                continue
                
            next_line = lines[i+1].strip()
            
            # --- ルール1: 床面積の結合 ---
            # 「床面積」が含まれる行の次が、数字や「X階」で始まる場合、結合する
            if "床面積" in line:
                # 次の行が数字、または「○階」で始まっているか？
                if re.match(r'^[\d０-９]+', next_line) or re.match(r'^[1-9１-９]階', next_line):
                    # 行を結合 (全角スペース区切り)
                    merged_line = lines[i].rstrip() + "　" + next_line
                    processed_lines.append(merged_line)
                    skip_next = True
                    continue

            # --- ルール2: 所在の結合 ---
            # 「所在」が含まれる行の次が、数字で始まっている（番地の続き）場合、結合する
            # 例: "所在 四街道市..." の次の行が "1520..."
            if "所在" in line:
                # 次の行が数字で始まっているか？ (全角半角問わず)
                # かつ、次の行が「家屋番号」などの別のヘッダーではないことを確認
                is_number_start = re.match(r'^[\d０-９]+', next_line)
                is_header = any(x in next_line for x in ["家屋番号", "地番", "地目", "種類", "構造", "床面積", "地積", "持分"])
                
                if is_number_start and not is_header:
                    # 番地っぽさを出すために、数字だけなら「番地」などを補完しても良いが、
                    # ここではシンプルに結合する
                    # ユーザー要望: "1520番地236" のようにしたい
                    
                    # もし次の行に「番」が含まれていなければ、「番地」を補完するロジック（オプション）
                    # 今回は単純結合 + 番地補完を試みる
                    if "番" not in next_line and "地" not in next_line:
                        # 数字だけの羅列なら「番地」を挟む？ -> リスクがあるので単純結合にする
                        pass
                    
                    merged_line = lines[i].rstrip() + "　" + next_line
                    processed_lines.append(merged_line)
                    skip_next = True
                    continue

            processed_lines.append(lines[i])

        return "\n".join(processed_lines)

    def _trim_whitespace(self, img: Image.Image) -> Image.Image:
        try:
            bg = Image.new(img.mode, img.size, (255, 255, 255))
            diff = ImageChops.difference(img, bg)
            diff = ImageChops.add(diff, diff, 2.0, -100)
            bbox = diff.getbbox()
            if bbox:
                return img.crop(bbox)
        except: pass
        return img

    def _invoke_ai_reasoning(self, input_text: str) -> WillDraftStructure:
        system_content = """
        あなたは熟練した行政書士です。提供された「遺産整理要旨」に基づき、公正証書遺言の条文案を作成してください。
        （中略: プロンプトは変更なし）
        出力は指定されたJSONスキーマに厳密に従ってください。
        """
        
        parser = PydanticOutputParser(pydantic_object=WillDraftStructure)
        
        prompt = ChatPromptTemplate.from_messages([
            SystemMessage(content=system_content),
            HumanMessagePromptTemplate.from_template(
                "以下の要旨データに基づき、遺言書ドラフトの【本文条項のみ】を作成してください。\n\n【要旨データ】\n{input_text}\n\n【出力形式】\n{format_instructions}"
            )
        ])
        
        chain = prompt | self.llm | parser
        return chain.invoke({
            "input_text": input_text,
            "format_instructions": parser.get_format_instructions()
        })

    def _set_jp_font(self, run, size_pt=12, is_bold=False):
        try:
            run.font.name = "MS Mincho"
            run.font.size = Pt(size_pt)
            run.font.bold = is_bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'MS Mincho')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'MS Mincho')
        except Exception:
            pass

    def _create_will_document(self, template_file: BytesIO, data: WillDraftStructure, registry_text: str = "") -> Document:
        """遺言書本体の作成（テンプレート追記モード）"""
        try:
            doc = Document(template_file)
        except Exception:
            doc = Document() 

        doc.add_paragraph("\n") 

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        timestamp = datetime.now().strftime('%Y年%m月%d日 ドラフト作成')
        self._set_jp_font(p_date.add_run(timestamp), size_pt=9)
        doc.add_paragraph("") 

        if not data.articles:
            doc.add_paragraph("※ 生成された条文データがありません。要旨の内容を確認してください。")
            return doc

        for article in data.articles:
            p_title = doc.add_paragraph()
            self._set_jp_font(p_title.add_run(f"{article.article_number}"), size_pt=12, is_bold=True)
            if article.title:
                self._set_jp_font(p_title.add_run(f"　（{article.title}）"), size_pt=12, is_bold=True)
            
            p_content = doc.add_paragraph()
            p_content.paragraph_format.first_line_indent = Mm(5)
            
            content_text = article.content if article.content else ""
            
            if "※要確認" in content_text:
                parts = content_text.split("（※要確認")
                self._set_jp_font(p_content.add_run(parts[0]), size_pt=12)
                if len(parts) > 1:
                    run_alert = p_content.add_run(f"（※要確認{parts[1]}")
                    self._set_jp_font(run_alert, size_pt=12, is_bold=True)
                    run_alert.font.color.rgb = RGBColor(255, 0, 0)
            else:
                self._set_jp_font(p_content.add_run(content_text), size_pt=12)
            
            doc.add_paragraph("")

        if data.supplementary_provisions:
            p_head = doc.add_paragraph()
            self._set_jp_font(p_head.add_run("（付言事項）"), size_pt=12, is_bold=True)
            p_body = doc.add_paragraph()
            p_body.paragraph_format.first_line_indent = Mm(5)
            self._set_jp_font(p_body.add_run(data.supplementary_provisions), size_pt=12)

        if registry_text:
            doc.add_page_break()
            p_ht = doc.add_paragraph()
            p_ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_jp_font(p_ht.add_run("【参考】不動産登記情報（テキストデータ）"), size_pt=14, is_bold=True)
            doc.add_paragraph("※公証人作成用の参考テキストです。\n")
            
            p_txt = doc.add_paragraph(registry_text)
            if p_txt.runs:
                self._set_jp_font(p_txt.runs[0], size_pt=10.5)
            else:
                self._set_jp_font(p_txt.add_run(registry_text), size_pt=10.5)

        return doc

    def _create_registry_document(self, registry_data: Dict[str, Any]) -> Document:
        """登記情報（別冊・画像のみ）の作成"""
        doc = Document()
        
        p_main = doc.add_paragraph()
        p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_jp_font(p_main.add_run("【別冊】不動産登記情報"), size_pt=20, is_bold=True)
        doc.add_paragraph("\n")
        
        images = registry_data.get("images", [])
        if images:
            for img_data in images:
                try:
                    img_data.seek(0)
                    doc.add_picture(img_data, width=Mm(170))
                    doc.add_paragraph("") 
                except Exception as e:
                    doc.add_paragraph(f"※画像エラー: {e}")
        
        return doc